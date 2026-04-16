"""
Template générique de lettre de recommandation professionnelle
Utiliser python-docx : pip install python-docx

Usage:
1. Copier ce fichier et renommer : generate_lettre_[prenom].py
2. Modifier les variables en haut (PERSONNE, POSTE, etc.)
3. Personnaliser les paragraphes
4. Exécuter : python generate_lettre_[prenom].py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ════════════════════════════════════════════════════════════════════════════
# CONFIGURATION - À PERSONNALISER
# ════════════════════════════════════════════════════════════════════════════

# Personne recommandée
PERSONNE_NOM = "Prénom Nom"
PERSONNE_PRENOM = "Prénom"

# Poste visé
POSTE_VISE = "Data Engineer senior ou Lead Data"

# Signataire (par défaut : Yann Ponaire Infopro Digital)
SIGNATAIRE_NOM = "Yann Ponaire"
SIGNATAIRE_TITRE = "Deputy CTO / Directeur Technique Data"
SIGNATAIRE_ENTREPRISE = "Infopro Digital"
SIGNATAIRE_DATE = "Mars 2026"

# Durée collaboration
DUREE_COLLABORATION = "deux ans"

# Compétences techniques (stack)
COMPETENCES_STACK = "Talend, Snowflake, Python, Airbyte, n8n"

# ════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DU DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

# ── HELPER FUNCTIONS ────────────────────────────────────────────────────────

def add_paragraph(doc, text="", bold=False, italic=False, font_name="Calibri",
                  font_size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    """Ajoute un paragraphe avec formatage uniforme"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return p

def add_mixed_paragraph(doc, segments, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=8):
    """Paragraphe avec styles mixtes (gras/italique dans le même paragraphe)
    segments = list of (text, bold, italic, font_name, font_size)
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    for text, bold, italic, font_name, font_size in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return p

# ── EN-TÊTE (aligné droite) ─────────────────────────────────────────────────
add_paragraph(doc, SIGNATAIRE_NOM, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
add_paragraph(doc, SIGNATAIRE_TITRE, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
add_paragraph(doc, SIGNATAIRE_ENTREPRISE, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
add_paragraph(doc, SIGNATAIRE_DATE, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=20)

# ── OBJET ───────────────────────────────────────────────────────────────────
add_mixed_paragraph(doc, [
    ("Objet : ", True, False, "Calibri", 11),
    (f"Lettre de recommandation — {PERSONNE_NOM}", False, False, "Calibri", 11)
], align=WD_ALIGN_PARAGRAPH.LEFT, space_after=16)

# ── FORMULE D'ADRESSE ───────────────────────────────────────────────────────
add_paragraph(doc, "À qui de droit,", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║ CONTENU À PERSONNALISER                                                   ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

# ── ACCROCHE (en gras) ──────────────────────────────────────────────────────
add_paragraph(doc, f"{PERSONNE_NOM} est le type de profil qu'on ne remplace pas.", bold=True, space_after=10)

# ── PARAGRAPHE 1 : Contexte & observation générale ──────────────────────────
add_mixed_paragraph(doc, [
    (f"En plus de {DUREE_COLLABORATION} de collaboration chez {SIGNATAIRE_ENTREPRISE}, "
     "je l'ai observé dans des situations où la plupart des ingénieurs auraient décroché le téléphone "
     "pour demander de l'aide. "
     f"{PERSONNE_PRENOM}, lui, revient quelques jours plus tard avec un projet structuré, "
     "une documentation propre — et la question : ", False, False, "Calibri", 11),
    ("« J'ai plus de taff, je fais quoi ? »", False, True, "Calibri", 11),
], space_after=4)
add_paragraph(doc, "Ce n'est pas une formule. C'est sa façon de travailler.", italic=True, space_after=12)

# ── PARAGRAPHE 2 : Compétences techniques ───────────────────────────────────
add_mixed_paragraph(doc, [
    ("Compétences techniques : ", True, False, "Calibri", 11),
    (f"{PERSONNE_PRENOM} maîtrise l'ensemble du stack data moderne — {COMPETENCES_STACK} — "
     "et s'approprie chaque nouvel outil en quelques jours sans formation préalable. "
     "[AJOUTER UN EXEMPLE CONCRET : projet réalisé, outil adopté seul, initiative technique notable]",
     False, False, "Calibri", 11),
], space_after=10)

# ── PARAGRAPHE 3 : Autonomie & soft skills ──────────────────────────────────
add_mixed_paragraph(doc, [
    ("Autonomie et fiabilité : ", True, False, "Calibri", 11),
    (f"Là où certains profils ont besoin de cadre et de suivi, {PERSONNE_PRENOM} a besoin de sujets. "
     "On peut lui confier un périmètre flou, un pipeline cassé ou une migration incertaine : "
     "il revient avec quelque chose de solide. Il ne gère pas l'urgence — "
     "il absorbe la complexité et la transforme en livrable.",
     False, False, "Calibri", 11),
], space_after=10)

# ── PARAGRAPHE 4 (Optionnel) : Attitude / anecdote marquante ────────────────
add_mixed_paragraph(doc, [
    ("Attitude : ", True, False, "Calibri", 11),
    (f"Ce qui distingue {PERSONNE_PRENOM} au-delà du technique, c'est l'appétit. "
     "Il demande du travail. Il aime réfléchir, concevoir, construire. "
     "Dans une équipe data, cette énergie est contagieuse et rare.",
     False, False, "Calibri", 11),
], space_after=12)

# ── CONCLUSION ──────────────────────────────────────────────────────────────
add_paragraph(
    doc,
    f"Je recommande {PERSONNE_NOM} sans réserve pour tout poste de {POSTE_VISE}. "
    "Toute équipe qui l'intègre gagnera en productivité, en fiabilité et en sérénité opérationnelle.",
    space_after=10
)
add_paragraph(doc, "Je reste disponible pour tout échange complémentaire.", space_after=24)

# ── FORMULE DE POLITESSE ────────────────────────────────────────────────────
add_paragraph(doc, "Cordialement,", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)

# ── SIGNATURE MANUSCRITE ────────────────────────────────────────────────────
p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(4)
p_sig.paragraph_format.space_after = Pt(2)
run_sig = p_sig.add_run(SIGNATAIRE_NOM.split()[0][0] + " " + SIGNATAIRE_NOM.split()[1])  # "Y Ponaire"
run_sig.font.name = "Segoe Script"
run_sig.font.size = Pt(18)
run_sig.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)  # bleu marine stylo

# ── TITRE SOUS SIGNATURE ────────────────────────────────────────────────────
add_paragraph(doc, SIGNATAIRE_NOM, bold=True, font_size=10, space_after=2)
add_paragraph(doc, f"{SIGNATAIRE_TITRE} — {SIGNATAIRE_ENTREPRISE}", font_size=10, space_after=0)

# ── SAVE ────────────────────────────────────────────────────────────────────
prenom_lower = PERSONNE_PRENOM.lower().replace(" ", "_")
nom_lower = PERSONNE_NOM.split()[-1].lower().replace(" ", "_")
output_path = f"lettre_recommandation_{prenom_lower}_{nom_lower}.docx"
doc.save(output_path)
print(f"✅ Fichier généré : {output_path}")
